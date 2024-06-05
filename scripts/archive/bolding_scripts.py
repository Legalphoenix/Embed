#update definitions in article 4 of gdpr
'''import chromadb
import re

# Initialize ChromaDB client
chroma_client = chromadb.HttpClient(host='localhost', port=8000)

# Fetch the legislation collection
legislation_collection_name = "legislation"
legislation_collection = chroma_client.get_collection(name=legislation_collection_name)

# Function to retrieve all IDs and full_preview_texts from a collection
def retrieve_full_preview_texts(collection):
    results = collection.get(include=["metadatas"])
    metadatas = results['metadatas']
    ids = results['ids']
    
    id_text_pairs = []
    for doc_id, metadata in zip(ids, metadatas):
        full_preview_text = metadata.get("full_preview_text", "No full_preview_text available")
        id_text_pairs.append((doc_id, full_preview_text))
    
    return id_text_pairs

# Function to update full_preview_text
def update_full_preview_text(collection, doc_id, full_preview_text):
    collection.update(
        ids=[doc_id],
        metadatas=[{"full_preview_text": full_preview_text}]
    )

# Function to add <b> tags around terms being defined in 'full_preview_text'
def bold_defined_terms(text):
    # Regex to find terms in single quotes following a number and a full stop
    pattern = r"(\d+\.\s*‘[^’]+’)"
    bolded_text = re.sub(pattern, lambda match: match.group(1).replace('‘', '<b>‘').replace('’', '’</b>'), text)
    return bolded_text

# Retrieve all IDs and full_preview_texts from legislation collection
legislation_id_text_pairs = retrieve_full_preview_texts(legislation_collection)

# Find all items with full_preview_text starting with "Article 4" or "<b>Article 4</b>" on the first line
matching_ids = []
for doc_id, full_preview_text in legislation_id_text_pairs:
    lines = full_preview_text.split('\n')
    if lines and (lines[0].strip() == "Article 4" or lines[0].strip() == "<b>Article 4</b>"):
        matching_ids.append((doc_id, full_preview_text))

# Check how many matches were found
if len(matching_ids) == 0:
    print("No document found with 'Article 4' on the first line.")
elif len(matching_ids) > 1:
    print("Multiple documents found with 'Article 4' on the first line. Document IDs:")
    for doc_id, _ in matching_ids:
        print(doc_id)
else:
    # Only one match found, proceed with updating
    doc_id, full_preview_text = matching_ids[0]
    modified_text = bold_defined_terms(full_preview_text)
    update_full_preview_text(legislation_collection, doc_id, modified_text)
    print(f"Updated full_preview_text for document ID: {doc_id}")

print("Script execution completed.")'''


#update with bold text for recitals collection
'''import chromadb
import re

# Initialize ChromaDB client
chroma_client = chromadb.HttpClient(host='localhost', port=8000)

# Fetch the recitals collection
recitals_collection_name = "recitals"
recitals_collection = chroma_client.get_collection(name=recitals_collection_name)

# Function to retrieve all IDs and full_preview_texts from a collection
def retrieve_full_preview_texts(collection):
    results = collection.get(include=["metadatas"])
    metadatas = results['metadatas']
    ids = results['ids']
    
    id_text_pairs = []
    for doc_id, metadata in zip(ids, metadatas):
        full_preview_text = metadata.get("full_preview_text", "No full_preview_text available")
        id_text_pairs.append((doc_id, full_preview_text))
    
    return id_text_pairs

# Function to update full_preview_text
def update_full_preview_text(collection, doc_id, full_preview_text):
    collection.update(
        ids=[doc_id],
        metadatas=[{"full_preview_text": full_preview_text}]
    )

# Function to add <b> tags around "Recital X" on the first line
def add_bold_tags_first_line(text):
    lines = text.split('\n')
    if len(lines) > 0 and re.match(r"^Recital \d+$", lines[0].strip()):
        lines[0] = f"<b>{lines[0]}</b>"
    return '\n'.join(lines)

# Retrieve all IDs and full_preview_texts from recitals collection
recitals_id_text_pairs = retrieve_full_preview_texts(recitals_collection)

# Process each text in recitals collection, modify it, and update the collection
for doc_id, full_preview_text in recitals_id_text_pairs:
    modified_text = add_bold_tags_first_line(full_preview_text)
    update_full_preview_text(recitals_collection, doc_id, modified_text)
    print(f"Updated full_preview_text for document ID: {doc_id} in recitals collection")

print("All full_preview_texts in recitals collection have been updated with <b> tags on the first line.")'''


#Update with bold text for legilsation
'''import chromadb
import re

# Initialize ChromaDB client
chroma_client = chromadb.HttpClient(host='localhost', port=8000)

# Fetch the collection
collection_name = "legislation"
collection = chroma_client.get_collection(name=collection_name)

# Function to retrieve all IDs and full_preview_texts from the collection
def retrieve_full_preview_texts(collection):
    results = collection.get(include=["metadatas"])
    metadatas = results['metadatas']
    ids = results['ids']
    
    id_text_pairs = []
    for doc_id, metadata in zip(ids, metadatas):
        full_preview_text = metadata.get("full_preview_text", "No full_preview_text available")
        id_text_pairs.append((doc_id, full_preview_text))
    
    return id_text_pairs

# Function to update full_preview_text
def update_full_preview_text(collection, doc_id, full_preview_text):
    collection.update(
        ids=[doc_id],
        metadatas=[{"full_preview_text": full_preview_text}]
    )

# Function to add <b> tags around "Article X" on the first line
def add_bold_tags_first_line(text):
    lines = text.split('\n')
    if lines and re.match(r"^Article d+$", lines[0]):
        lines[0] = f"<b>{lines[0]}</b>"
    return '\n'.join(lines)

# Retrieve all IDs and full_preview_texts
id_text_pairs = retrieve_full_preview_texts(collection)

# Process each text, modify it, and update the collection
for doc_id, full_preview_text in id_text_pairs:
    modified_text = add_bold_tags_first_line(full_preview_text)
    update_full_preview_text(collection, doc_id, modified_text)
    print(f"Updated full_preview_text for document ID: {doc_id}")

print("All full_preview_texts have been updated with <b> tags on the first line.")'''


'''import chromadb
import docx

# Initialize ChromaDB client
chroma_client = chromadb.HttpClient(host='localhost', port=8000)

# Fetch the collection
collection_name = "legislation"
collection = chroma_client.get_collection(name=collection_name)

# Function to retrieve all IDs and full_preview_texts from the collection
def retrieve_full_preview_texts(collection):
    results = collection.get(include=["metadatas"])
    metadatas = results['metadatas']
    ids = results['ids']
    
    id_text_pairs = []
    for doc_id, metadata in zip(ids, metadatas):
        full_preview_text = metadata.get("full_preview_text", "No full_preview_text available")
        id_text_pairs.append((doc_id, full_preview_text))
    
    return id_text_pairs

# Retrieve all IDs and full_preview_texts
id_text_pairs = retrieve_full_preview_texts(collection)

# Create a new document and add all IDs and full_preview_texts
doc = docx.Document()
for doc_id, full_preview_text in id_text_pairs:
    doc.add_paragraph(f"ID: {doc_id}")
    doc.add_paragraph(full_preview_text)
    doc.add_paragraph("\n")  # Add a newline for separation

# Save the document
output_path = '/Users/volumental/Desktop/Embed-search/legislation_full_preview_backup.docx'
doc.save(output_path)

print(f"All IDs and their full_preview_texts have been compiled into a single document at {output_path}.")'''
