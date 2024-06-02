#checkrecitals.py
import docx
from chromadb import HttpClient, Settings
import os

# Initialize ChromaDB client
chroma_client = HttpClient(
    host="localhost",
    port="8000",
    ssl=False,
    headers={},
    settings=Settings()
)

# Get the recitals collection
collection_recitals = chroma_client.get_collection(name="parent_court_cases")

# Function to retrieve all chunk texts from the recitals collection
def retrieve_chunk_texts(collection):
    chunk_texts = []
    results = collection.get(include=["documents"])
    for document in results['documents']:
        chunk_texts.append(document)
    return chunk_texts

# Retrieve all chunk texts
chunk_texts = retrieve_chunk_texts(collection_recitals)

# Create a new document and add all chunk texts
doc = docx.Document()
for chunk_text in chunk_texts:
    doc.add_paragraph(chunk_text)
    doc.add_paragraph("\n")  # Add a newline for separation

# Save the document
output_path = '/Users/volumental/Desktop/Embed-search/compiled_cases_document.docx'
doc.save(output_path)

print(f"All chunk texts have been compiled into a single document at {output_path}.")
