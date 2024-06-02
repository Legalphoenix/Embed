#Recitals.py
import docx
from docx import Document
import re
import os

# Load the document
doc = Document('input_document.docx')

# Initialize variables
recitals = []
current_recital = ""
recital_number = 1

# Regex pattern to identify recital headings
pattern = re.compile(r"^Recital (\d+)")

# Function to save recitals into separate documents
def save_recital_as_docx(recital_text, number):
    doc = Document()
    doc.add_paragraph(recital_text)
    doc.save(f"Recital_{number}.docx")

# Iterate over paragraphs and split by Recital headings
for paragraph in doc.paragraphs:
    if pattern.match(paragraph.text):
        if current_recital:
            recitals.append(current_recital.strip())
        current_recital = paragraph.text
    else:
        current_recital += "\n" + paragraph.text

# Add the last recital
if current_recital:
    recitals.append(current_recital.strip())

# Output the first 20 recitals
output_folder = "recitals_output"
os.makedirs(output_folder, exist_ok=True)
os.chdir(output_folder)

for i, recital in enumerate(recitals[:20], 1):
    save_recital_as_docx(recital, i)

print(f"Saved the first 20 recitals as individual documents in '{output_folder}' directory.")
