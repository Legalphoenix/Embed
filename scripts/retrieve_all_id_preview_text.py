import chromadb
import docx

# Initialize ChromaDB client
chroma_client = chromadb.HttpClient(host='localhost', port=8000)

# Fetch the collection
collection_name = "parent_court_cases"
collection = chroma_client.get_collection(name=collection_name)

# Function to retrieve all IDs and full_preview_texts from the collection
def retrieve_full_preview_texts(collection):
    results = collection.get(include=["metadatas"])
    metadatas = results['metadatas']
    ids = results['ids']
    
    id_text_pairs = []
    for doc_id, metadata in zip(ids, metadatas):
        full_preview_text = metadata.get("chunk_text", "No full_preview_text available")
        id_text_pairs.append((doc_id, full_preview_text))
    
    return id_text_pairs

# Retrieve all IDs and full_preview_texts
id_text_pairs = retrieve_full_preview_texts(collection)

# Create a new document and add all IDs and full_preview_texts
doc = docx.Document()
for doc_id, full_preview_text in id_text_pairs:
    doc.add_paragraph(f"ID: {doc_id}")
    doc.add_paragraph(full_preview_text)
    doc.add_paragraph("\n")  # Add a newline for separation

# Save the document
output_path = '/Users/volumental/Desktop/Embed-search/legislation_full_preview_backup.docx'
doc.save(output_path)

print(f"All IDs and their full_preview_texts have been compiled into a single document at {output_path}.")